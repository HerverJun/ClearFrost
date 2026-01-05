from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_admin_manual():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei UI'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    
    # Title
    title = doc.add_heading('视觉检测系统管理员手册', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('GreeVision Rebirth V1.0 | 后台配置指南')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    
    doc.add_paragraph()
    
    # ========== Section 1 ==========
    doc.add_heading('一、进入后台设置', level=1)
    
    doc.add_heading('1.1 入口位置', level=2)
    doc.add_paragraph('点击软件界面右上角的 ⚙️ 齿轮图标，即可进入后台设置。')
    
    doc.add_heading('1.2 密码验证', level=2)
    doc.add_paragraph('为防止现场操作员误操作，后台设置需要输入管理员密码。')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('默认密码：').bold = True
    p.add_run('888888')
    doc.add_paragraph()
    warn = doc.add_paragraph()
    warn.add_run('⚠️ 安全提醒：').bold = True
    warn.add_run('如需修改密码，请联系软件开发人员在代码中进行更改。建议在正式投产前更换为自定义密码。')
    
    # ========== Section 2 ==========
    doc.add_heading('二、后台参数说明', level=1)
    
    doc.add_paragraph('后台设置界面分为以下几个配置区域：')
    
    # --- 2.1 ---
    doc.add_heading('2.1 存储路径', level=2)
    
    p = doc.add_paragraph()
    p.add_run('参数名称：').bold = True
    p.add_run('存储路径')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('作用：').bold = True
    p.add_run('指定检测图片和日志文件的保存位置。')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('修改后影响：').bold = True
    doc.add_paragraph('• 所有新产生的NG图片将保存到新路径下')
    doc.add_paragraph('• 日志文件也会写入新路径')
    doc.add_paragraph('• 旧路径下的历史数据不会自动迁移，如需保留请手动复制')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('建议：').bold = True
    p.add_run('选择磁盘空间充足的分区（如D盘），避免C盘空间不足影响系统运行。')
    
    # --- 2.2 ---
    doc.add_heading('2.2 PLC配置', level=2)
    
    doc.add_paragraph('此区域用于配置与PLC的通讯参数。')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('IP地址：').bold = True
    p.add_run('PLC的网络地址，例如 192.168.1.100')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('端口：').bold = True
    p.add_run('PLC的通讯端口号，常见为 502（Modbus TCP）')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('触发地址：').bold = True
    p.add_run('PLC发送"开始检测"信号的寄存器地址')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('结果地址：').bold = True
    p.add_run('软件向PLC反馈检测结果的寄存器地址')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('修改后影响：').bold = True
    doc.add_paragraph('• 修改IP/端口后，需点击 [连接PLC] 按钮重新建立连接')
    doc.add_paragraph('• 地址配置错误会导致PLC无法触发检测或无法接收结果')
    doc.add_paragraph('• 修改前务必与电气工程师确认正确的点位信息')
    
    # --- 2.3 ---
    doc.add_heading('2.3 相机配置', level=2)
    
    doc.add_paragraph('此区域用于配置工业相机参数。')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('显示名称：').bold = True
    p.add_run('相机在界面上显示的名称，仅用于标识，不影响实际功能')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('序列号：').bold = True
    p.add_run('相机的唯一序列号（SN），用于指定连接哪台相机。可在相机机身标签或海康MVS软件中查看')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('曝光 (ms)：').bold = True
    p.add_run('相机曝光时间，单位为毫秒。数值越大画面越亮，但可能产生运动模糊')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('增益：').bold = True
    p.add_run('信号放大倍数。增益越高画面越亮，但噪点也会增加')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('修改后影响：').bold = True
    doc.add_paragraph('• 曝光/增益的调整会立即影响画面亮度')
    doc.add_paragraph('• 画面过亮或过暗都会影响AI识别准确率')
    doc.add_paragraph('• 建议在现场光照稳定后，通过实际测试确定最佳参数')
    
    # --- 2.4 ---
    doc.add_heading('2.4 高级设置（检测逻辑）', level=2)
    
    doc.add_paragraph('此区域用于配置AI检测的判定逻辑。')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('目标标签：').bold = True
    p.add_run('需要检测的目标名称，例如 "screw"（螺钉）、"remote"（遥控器）等。必须与AI模型训练时的类别名称一致')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('目标数量：').bold = True
    p.add_run('判定为"合格"所需的目标数量。例如设置为4，则检测到4个或以上目标时判定为合格')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('启用GPU加速：').bold = True
    p.add_run('是否使用显卡进行AI推理。启用后速度更快，但需要电脑配备NVIDIA显卡并安装CUDA驱动')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('修改后影响：').bold = True
    doc.add_paragraph('• 目标标签配置错误会导致检测失败或始终判定不合格')
    doc.add_paragraph('• 目标数量设置不当会影响良率统计（设太低则漏检，设太高则误判）')
    doc.add_paragraph('• GPU加速开关修改后需重启软件生效')
    
    # ========== Section 3 ==========
    doc.add_heading('三、界面上的阈值滑块', level=1)
    
    doc.add_paragraph('除了后台设置外，主界面上还有两个可实时调整的滑块：')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('置信度 (Confidence)：').bold = True
    p.add_run('AI识别结果的可信程度阈值，范围0.10~0.95。数值越高要求越严格，可能漏检；数值越低要求越宽松，可能误检')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('IOU阈值：').bold = True
    p.add_run('用于过滤重叠检测框的阈值。一般保持默认0.30即可，非专业人员不建议修改')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('调整建议：').bold = True
    doc.add_paragraph('• 如果频繁出现漏检（明明有螺钉却判不合格），可尝试降低置信度')
    doc.add_paragraph('• 如果频繁出现误检（没有螺钉却判合格），可尝试提高置信度')
    doc.add_paragraph('• 每次调整后请观察多个产品的检测情况，不要频繁大幅修改')
    
    # ========== Section 4 ==========
    doc.add_heading('四、配置修改流程（推荐）', level=1)
    
    doc.add_paragraph('为确保配置修改不影响正常生产，建议按以下步骤操作：')
    doc.add_paragraph()
    
    doc.add_paragraph('1. 在非生产时段（如换班间隙）进行修改', style='List Number')
    doc.add_paragraph('2. 修改前记录当前参数值（拍照或笔记）', style='List Number')
    doc.add_paragraph('3. 一次只修改一个参数，便于定位问题', style='List Number')
    doc.add_paragraph('4. 修改后进行至少5次手动检测验证', style='List Number')
    doc.add_paragraph('5. 确认无误后再恢复自动生产模式', style='List Number')
    doc.add_paragraph('6. 如出现异常，立即恢复原参数值', style='List Number')
    
    # ========== Section 5 ==========
    doc.add_heading('五、常见配置问题', level=1)
    
    doc.add_heading('Q1：修改参数后检测一直不合格', level=2)
    doc.add_paragraph('首先检查"目标标签"是否填写正确（区分大小写）。其次检查"目标数量"是否设置过高。最后尝试降低"置信度"滑块。')
    
    doc.add_heading('Q2：PLC无法连接', level=2)
    doc.add_paragraph('确认IP地址和端口填写正确。检查网线是否连接正常。在电脑上ping一下PLC的IP，看是否能通。如仍无法解决，联系电气工程师。')
    
    doc.add_heading('Q3：相机画面过亮/过暗', level=2)
    doc.add_paragraph('调整"曝光"和"增益"参数。曝光影响整体亮度，增益影响信号放大。建议先调曝光，不够再调增益。避免增益设置过高导致噪点。')
    
    doc.add_heading('Q4：如何部署到新工位', level=2)
    doc.add_paragraph('将整个软件文件夹复制到新电脑。根据新工位的硬件，修改相机序列号、PLC地址等参数。如检测目标不同，需更换AI模型文件并修改目标标签。')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('文档版本：V1.0 | 更新日期：2025-12-23 | 仅限内部使用')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    
    # Save
    output_path = r'c:\Users\11234\Desktop\W5电机螺钉检测 -Rebirth\视觉检测系统管理员手册.docx'
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')

if __name__ == '__main__':
    create_admin_manual()
