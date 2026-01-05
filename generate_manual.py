from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_manual():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei UI'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    
    # Title
    title = doc.add_heading('视觉检测系统操作手册', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('GreeVision Rebirth V1.0 | 现场操作指南')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    
    doc.add_paragraph()
    
    # Section 1
    doc.add_heading('一、界面总览', level=1)
    p = doc.add_paragraph('软件界面分为四个主要区域：')
    
    # Layout description
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Row 1 - Header
    cell = table.cell(0, 0)
    cell.merge(table.cell(0, 2))
    cell.text = '顶部栏：软件名称 + 功能按钮（日志/图库/设置/窗口控制）'
    set_cell_shading(cell, 'E2E8F0')
    
    # Row 2
    table.cell(1, 0).text = '相机画面区域\n（主显示区）'
    table.cell(1, 1).text = '状态指示灯\n今日统计\n控制面板'
    table.cell(1, 2).text = '检测记录\n（实时日志）'
    
    # Row 3
    table.cell(2, 0).text = ''
    table.cell(2, 1).text = ''
    table.cell(2, 2).text = '系统日志\n（状态信息）'
    
    doc.add_paragraph()
    
    # Section 2
    doc.add_heading('二、开机启动', level=1)
    
    doc.add_heading('第一步：启动软件', level=2)
    doc.add_paragraph('双击桌面上的 "视觉检测系统" 图标，等待软件加载完成。')
    
    doc.add_heading('第二步：检查连接状态', level=2)
    doc.add_paragraph('观察界面中间上方的状态指示灯：')
    
    # Status table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '指示灯'
    table.cell(0, 1).text = '绿色亮起'
    table.cell(0, 2).text = '灰色熄灭'
    set_cell_shading(table.cell(0, 0), 'E2E8F0')
    set_cell_shading(table.cell(0, 1), 'E2E8F0')
    set_cell_shading(table.cell(0, 2), 'E2E8F0')
    
    table.cell(1, 0).text = '相机通讯'
    table.cell(1, 1).text = '✅ 相机已连接'
    table.cell(1, 2).text = '❌ 相机未连接'
    
    table.cell(2, 0).text = 'PLC通讯'
    table.cell(2, 1).text = '✅ PLC已连接'
    table.cell(2, 2).text = '❌ PLC未连接'
    
    doc.add_paragraph()
    p = doc.add_paragraph('如果指示灯为灰色：')
    doc.add_paragraph('1. 点击 [查找相机] 按钮', style='List Number')
    doc.add_paragraph('2. 点击 [打开相机] 按钮', style='List Number')
    doc.add_paragraph('3. 点击 [连接PLC] 按钮', style='List Number')
    doc.add_paragraph('4. 如仍无法连接，请联系技术人员', style='List Number')
    
    doc.add_heading('第三步：确认画面显示', level=2)
    doc.add_paragraph('左侧大屏幕应显示相机正在拍摄的画面。如果显示"等待信号"，说明相机未正常工作。')
    
    # Section 3
    doc.add_heading('三、日常操作', level=1)
    
    doc.add_heading('3.1 自动检测（正常生产模式）', level=2)
    doc.add_paragraph('当PLC发送触发信号时，系统会自动完成以下流程：')
    doc.add_paragraph('1. 自动拍照', style='List Number')
    doc.add_paragraph('2. AI分析图片', style='List Number')
    doc.add_paragraph('3. 显示检测结果（合格 或 不合格）', style='List Number')
    doc.add_paragraph('4. 将结果反馈给PLC', style='List Number')
    doc.add_paragraph('5. 统计数据自动更新', style='List Number')
    
    tip = doc.add_paragraph()
    tip.add_run('💡 提示：').bold = True
    tip.add_run('操作员无需干预，系统全自动运行。')
    
    doc.add_heading('3.2 手动检测（调试/抽检模式）', level=2)
    doc.add_paragraph('如需手动触发一次检测（例如换班抽检），请点击 [手动检测] 按钮。')
    
    doc.add_heading('3.3 手动放行', level=2)
    doc.add_paragraph('若产品被误判为不合格，确认无问题后，可点击 [手动放行] 按钮强制放行。')
    warn = doc.add_paragraph()
    warn.add_run('⚠️ 谨慎使用：').bold = True
    warn.add_run('放行记录会被系统记录。')
    
    doc.add_heading('3.4 查看今日统计', level=2)
    doc.add_paragraph('界面右侧中间区域显示当天的检测统计：')
    doc.add_paragraph('• 总计：今日检测总数量')
    doc.add_paragraph('• 合格：通过检测的数量（绿色）')
    doc.add_paragraph('• 不合格：未通过检测的数量（红色）')
    
    # Section 4
    doc.add_heading('四、查看历史记录', level=1)
    
    doc.add_heading('4.1 查看检测日志', level=2)
    doc.add_paragraph('1. 点击顶部栏的 📄文档图标（检测日志按钮）', style='List Number')
    doc.add_paragraph('2. 弹出窗口显示历史检测记录', style='List Number')
    doc.add_paragraph('3. 每条记录包含：检测时间、结果、详情', style='List Number')
    doc.add_paragraph('4. 点击 [刷新] 按钮可更新最新数据', style='List Number')
    
    doc.add_heading('4.2 查看不合格图片', level=2)
    doc.add_paragraph('1. 点击顶部栏的 🖼️图片图标（图片库按钮）', style='List Number')
    doc.add_paragraph('2. 左侧选择日期', style='List Number')
    doc.add_paragraph('3. 上方选择小时', style='List Number')
    doc.add_paragraph('4. 点击缩略图可放大查看', style='List Number')
    
    # Section 5
    doc.add_heading('五、常见问题处理', level=1)
    
    doc.add_heading('问题1：相机画面黑屏/无图像', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '检查项'
    table.cell(0, 1).text = '处理方法'
    set_cell_shading(table.cell(0, 0), 'E2E8F0')
    set_cell_shading(table.cell(0, 1), 'E2E8F0')
    table.cell(1, 0).text = '相机通讯指示灯'
    table.cell(1, 1).text = '若为灰色，点击 [查找相机] → [打开相机]'
    table.cell(2, 0).text = '相机电源'
    table.cell(2, 1).text = '检查相机电源线是否松动'
    table.cell(3, 0).text = '网线连接'
    table.cell(3, 1).text = '检查相机网线是否插紧'
    
    doc.add_paragraph()
    
    doc.add_heading('问题2：PLC不触发检测', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '检查项'
    table.cell(0, 1).text = '处理方法'
    set_cell_shading(table.cell(0, 0), 'E2E8F0')
    set_cell_shading(table.cell(0, 1), 'E2E8F0')
    table.cell(1, 0).text = 'PLC通讯指示灯'
    table.cell(1, 1).text = '若为灰色，点击 [连接PLC]'
    table.cell(2, 0).text = 'PLC运行状态'
    table.cell(2, 1).text = '检查PLC是否正常运行'
    table.cell(3, 0).text = '信号地址'
    table.cell(3, 1).text = '联系技术人员确认PLC配置'
    
    doc.add_paragraph()
    
    doc.add_heading('问题3：检测结果一直不合格', level=2)
    doc.add_paragraph('可能原因：')
    doc.add_paragraph('1. 产品确实有缺陷 — 正常情况', style='List Number')
    doc.add_paragraph('2. 光源异常 — 检查光源是否正常亮起', style='List Number')
    doc.add_paragraph('3. 相机脏污 — 轻轻擦拭相机镜头', style='List Number')
    doc.add_paragraph('4. 阈值设置不当 — 联系技术人员调整置信度', style='List Number')
    
    doc.add_heading('问题4：软件卡顿/无响应', level=2)
    doc.add_paragraph('1. 等待30秒，看是否自动恢复', style='List Number')
    doc.add_paragraph('2. 如无响应，点击右上角 [退出程序] 按钮关闭软件', style='List Number')
    doc.add_paragraph('3. 重新启动软件', style='List Number')
    doc.add_paragraph('4. 如反复出现，联系技术人员', style='List Number')
    
    # Section 6
    doc.add_heading('六、窗口控制按钮', level=1)
    doc.add_paragraph('顶部栏右侧有三个窗口控制按钮：')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '按钮'
    table.cell(0, 1).text = '功能'
    set_cell_shading(table.cell(0, 0), 'E2E8F0')
    set_cell_shading(table.cell(0, 1), 'E2E8F0')
    table.cell(1, 0).text = '➖'
    table.cell(1, 1).text = '最小化窗口'
    table.cell(2, 0).text = '🔲'
    table.cell(2, 1).text = '最大化/还原窗口'
    table.cell(3, 0).text = '🚪'
    table.cell(3, 1).text = '退出程序（会弹出确认框）'
    
    # Section 7
    doc.add_heading('七、注意事项', level=1)
    doc.add_paragraph('❌ 请勿随意修改设置：设置界面需要管理员密码，普通操作无需进入')
    doc.add_paragraph('❌ 请勿遮挡相机：确保相机能正常拍到产品')
    doc.add_paragraph('❌ 请勿关闭软件：生产期间请保持软件运行')
    doc.add_paragraph('✅ 定期清洁镜头：每班次开始前，用干净软布轻擦相机镜头')
    doc.add_paragraph('✅ 异常及时上报：发现任何异常情况，请及时通知技术人员')
    
    # Section 8
    doc.add_heading('八、紧急联系', level=1)
    doc.add_paragraph('如遇无法解决的问题，请联系：')
    doc.add_paragraph('• 现场技术员：[填写姓名/电话]')
    doc.add_paragraph('• 系统维护：[填写姓名/电话]')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('文档版本：V1.0 | 更新日期：2025-12-23')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    
    # Save
    output_path = r'c:\Users\11234\Desktop\W5电机螺钉检测 -Rebirth\视觉检测系统操作手册.docx'
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')

if __name__ == '__main__':
    create_manual()
