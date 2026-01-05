from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_full_manual():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei UI'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei UI')
    
    # Cover Page
    for i in range(5): doc.add_paragraph()
    
    title = doc.add_heading('视觉检测系统\n用户使用与维护手册', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph('GreeVision Rebirth V1.0 | 究极详尽版')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    for i in range(8): doc.add_paragraph()
    
    info = doc.add_paragraph('适用对象：现场操作员、设备管理员、电气工程师\n发布日期：2025年12月23日')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(12)
    info.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    
    doc.add_page_break()
    
    # Table of Contents (Simulated)
    doc.add_heading('目录', level=1)
    doc.add_paragraph('第一部分：现场操作指南 ........................................................... 3')
    doc.add_paragraph('    1.1 界面总览 ............................................................................ 3')
    doc.add_paragraph('    1.2 开机与启动 ........................................................................ 4')
    doc.add_paragraph('    1.3日常操作流程 ...................................................................... 5')
    doc.add_paragraph('    1.4 异常处理与注意事项 ......................................................... 6')
    doc.add_paragraph('第二部分：后台配置与管理 ..................................................... 7')
    doc.add_paragraph('    2.1 进入后台与权限 ................................................................. 7')
    doc.add_paragraph('    2.2 核心参数详解 ..................................................................... 8')
    doc.add_paragraph('    2.3 AI判定逻辑调整 ................................................................. 10')
    doc.add_paragraph('    2.4 系统维护建议 ..................................................................... 11')
    
    doc.add_page_break()
    
    # ================= PART 1 =================
    part1 = doc.add_heading('第一部分：现场操作指南', level=1)
    part1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_heading('1.1 界面总览', level=2)
    p = doc.add_paragraph('软件主界面设计简洁直观，分为四大功能区：')
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '1. 监控显示区（左侧）'
    table.cell(0, 1).text = '2. 状态与控制区（中上）'
    table.cell(1, 0).text = '3. 数据统计区（中下）'
    table.cell(1, 1).text = '4. 日志记录区（右侧）'
    set_cell_shading(table.cell(0, 0), 'F1F5F9')
    set_cell_shading(table.cell(0, 1), 'F1F5F9')
    set_cell_shading(table.cell(1, 0), 'F1F5F9')
    set_cell_shading(table.cell(1, 1), 'F1F5F9')
    
    doc.add_paragraph()
    doc.add_paragraph('• 监控显示区：实时显示经过AI处理的画面，包含检测框和结果标签。')
    doc.add_paragraph('• 状态指示灯：绿色代表正常连接，灰色代表断开。')
    doc.add_paragraph('• 统计数据：实时统计当班次的合格/不合格数量，重启软件后清零。')
    
    doc.add_heading('1.2 开机与启动', level=2)
    doc.add_paragraph('标准开机流程：')
    doc.add_paragraph('1. 开启工控机电源，进入Windows桌面。', style='List Number')
    doc.add_paragraph('2. 双击桌面图标 "视觉检测系统"。', style='List Number')
    doc.add_paragraph('3. 等待约5-10秒，软件界面完全显示。', style='List Number')
    doc.add_paragraph('4. 检查顶部指示灯：', style='List Number')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '指示灯'
    table.cell(0, 1).text = '正常状态'
    table.cell(0, 2).text = '异常状态'
    set_cell_shading(table.cell(0, 0), 'E2E8F0')
    
    table.cell(1, 0).text = '相机通讯'
    table.cell(1, 1).text = '🟢 绿色'
    table.cell(1, 2).text = '⚪ 灰色（需点击"打开相机"）'
    
    table.cell(2, 0).text = 'PLC通讯'
    table.cell(2, 1).text = '🟢 绿色'
    table.cell(2, 2).text = '⚪ 灰色（需点击"连接PLC"）'
    
    doc.add_heading('1.3 日常操作流程', level=2)
    doc.add_paragraph('A. 自动模式（推荐）', style='Heading 3')
    doc.add_paragraph('• 系统默认处于自动模式。')
    doc.add_paragraph('• 当流水线传感器触发信号时，系统自动拍照并判断。')
    doc.add_paragraph('• 操作员只需关注屏幕上的大字结果：')
    doc.add_paragraph('    ✅ PASS (绿色)：合格，流水线继续运行。')
    doc.add_paragraph('    ❌ FAIL (红色)：不合格，流水线报警/停机。')
    
    doc.add_paragraph('B. 手动干预', style='Heading 3')
    doc.add_paragraph('• 手动检测：在无料或调试时，点击 [手动检测] 模拟一次触发。')
    doc.add_paragraph('• 手动放行：若系统误判（实际产品合格），点击 [手动放行] 发送OK信号给PLC。')
    
    doc.add_heading('1.4 异常处理', level=2)
    doc.add_paragraph('Q: 画面全黑？')
    doc.add_paragraph('A: 检查相机镜头盖是否取下，光源是否开启。点击 [查找相机] 尝试重连。')
    doc.add_paragraph()
    doc.add_paragraph('Q: 一直判定不合格？')
    doc.add_paragraph('A: 检查镜头是否脏污。如果产品没问题，可能是从后台设置的阈值过高，需通知管理员调整。')
    
    doc.add_page_break()
    
    # ================= PART 2 =================
    part2 = doc.add_heading('第二部分：后台配置与管理', level=1)
    part2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph('本部分内容仅限管理员或技术人员操作。')
    
    doc.add_heading('2.1 进入后台', level=2)
    doc.add_paragraph('入口：点击软件右上角的 ⚙️ 图标。')
    doc.add_paragraph('密码：默认密码为 888888。')
    
    doc.add_heading('2.2 核心参数详解', level=2)
    
    doc.add_paragraph('A. 存储配置', style='Heading 3')
    doc.add_paragraph('• 存储路径：建议设置在非系统盘（如 D:\\Data），防止占满C盘导致系统崩溃。')
    doc.add_paragraph('• 注意：修改路径后，之前的历史图片不会自动迁移。')
    
    doc.add_paragraph('B. PLC通讯配置', style='Heading 3')
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '参数'
    table.cell(0, 1).text = '说明'
    set_cell_shading(table.cell(0, 0), 'E2E8F0')
    
    table.cell(1, 0).text = 'IP地址'
    table.cell(1, 1).text = 'PLC的固定IP，需与工控机在同一网段。'
    table.cell(2, 0).text = '端口'
    table.cell(2, 1).text = '通常为 502 (Modbus TCP) 或自定义端口。'
    table.cell(3, 0).text = '触发地址'
    table.cell(3, 1).text = 'PLC写入"1"触发拍照的寄存器地址。'
    table.cell(4, 0).text = '结果地址'
    table.cell(4, 1).text = '软件写入结果（1=OK, 2=NG）的寄存器地址。'
    
    doc.add_paragraph('C. 相机参数配置', style='Heading 3')
    doc.add_paragraph('• 序列号 (SN)：必须与实际连接的相机一致，否则无法打开相机。')
    doc.add_paragraph('• 曝光 (Exposure)：控制画面亮度。曝光过低画面黑，过高画面白且有拖影。')
    doc.add_paragraph('• 增益 (Gain)：辅助提亮。建议优先调曝光，最后调增益以减少噪点。')
    
    doc.add_heading('2.3 AI判定逻辑调整', level=2)
    doc.add_paragraph('这是V1.0版本的核心升级功能，支持逻辑热更。')
    doc.add_paragraph()
    doc.add_paragraph('1. 目标标签 (Target Label)：')
    doc.add_paragraph('   必须与模型训练时的标签一致（如 "screw"）。填错将导致系统"视而不见"。')
    doc.add_paragraph()
    doc.add_paragraph('2. 目标数量 (Target Count)：')
    doc.add_paragraph('   判定合格所需的最少数量。')
    doc.add_paragraph('   示例：一个电机上需要4颗螺钉，则设为4。少于4颗判NG，多于4颗判OK。')
    doc.add_paragraph()
    doc.add_paragraph('3. 置信度阈值 (界面滑块)：')
    doc.add_paragraph('   • 建议值：0.50 - 0.70')
    doc.add_paragraph('   • 现象：如果经常漏检（有螺钉没认出），调低该值。')
    doc.add_paragraph('   • 现象：如果经常误检（把影子认成螺钉），调高该值。')
    
    doc.add_heading('2.4 系统维护建议', level=2)
    doc.add_paragraph('1. 定期清理磁盘：虽然软件会自动分类存储，但建议每季度手动备份并清理一次 D:\\Data 下的旧图片。', style='List Bullet')
    doc.add_paragraph('2. 备份配置文件：软件根目录下的 AppConfig.json 存储了所有设置，建议备份该文件。', style='List Bullet')
    doc.add_paragraph('3. 严禁随意修改文件名：不要手动修改 exe 文件名或移动 html 文件夹位置，否则会导致程序无法运行。', style='List Bullet')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('本文档由系统自动生成 | 最终解释权归研发部门所有')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    
    # Save
    output_path = r'c:\Users\11234\Desktop\W5电机螺钉检测 -Rebirth\视觉检测系统完整手册.docx'
    doc.save(output_path)
    print(f'完整版手册已生成: {output_path}')

if __name__ == '__main__':
    create_full_manual()
